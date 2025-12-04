"""
Word 파일 생성 유틸리티 (카드형 디자인 + Tavily 추천 아이템)
"""

import os
import re
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def add_hyperlink(paragraph, text: str, url: str, color: str = "0066CC"):
    """Word 문단에 클릭 가능한 하이퍼링크 추가"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    hyperlink.append(r)
    paragraph._p.append(hyperlink)
    return hyperlink


def clean_text(text: Any) -> str:
    if not isinstance(text, str):
        return str(text)
    text = re.sub(r"\s*\n\s*", " ", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def normalize_learning_items(items: Any) -> List[str]:
    if not isinstance(items, list):
        return [clean_text(items)]

    cleaned = []
    for it in items:
        if not isinstance(it, str):
            continue
        t = it.strip()
        if t.startswith(("•", "-", "*")):
            t = t[1:].strip()
        cleaned.append(t)

    if len(cleaned) > 1 and all(len(x) == 1 for x in cleaned):
        full_text = "".join(cleaned)
        sentences = re.split(r'(?<=[\.\?\!,])\s*', full_text)
        return [s.strip() for s in sentences if s.strip()]

    return cleaned


def add_card_header(doc: Document, text: str, color: RGBColor = RGBColor(0, 102, 204)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    line = doc.add_paragraph("━" * 35)
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_after = Pt(14)
    return p


def select_recommended_items(breakdown: list, category: str):
    if not breakdown:
        return []
    if category in ("Academic / STEM", "Career / Tech Skills"):
        primary = [item for item in breakdown if item.get("type") == "books"]
        secondary = [item for item in breakdown if item.get("type") != "books"]
    else:
        primary = [item for item in breakdown if item.get("type") != "books"]
        secondary = [item for item in breakdown if item.get("type") == "books"]
    return (primary + secondary)[:3]


def save_learning_guide_to_word(guide: Dict[str, Any], filename: Optional[str] = None) -> Optional[str]:
    if "error" in guide:
        print("❌ 가이드가 생성되지 않아 워드 파일을 만들 수 없습니다.")
        return None

    if filename is None:
        topic = guide.get("topic", "학습가이드").replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{topic}_학습가이드_{timestamp}.docx"

    doc = Document()
    set_document_style(doc)

    # 표지 카드
    title = doc.add_paragraph()
    title_run = title.add_run(f"📘 {guide.get('topic', '학습 가이드')}")
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 80, 160)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(guide.get("category", ""))
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("━" * 35).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # 학습 개요 카드
    add_card_header(doc, "📅 학습 개요")
    overview = [
        f"• 학습 기간: {guide.get('start_date', 'N/A')} ~ {guide.get('end_date', 'N/A')}",
        f"• 총 학습 일수: {guide.get('total_duration_days', 'N/A')}일",
    ]
    for line in overview:
        doc.add_paragraph(line)
    divider = doc.add_paragraph("━" * 35)
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # 비용 카드 + 추천 Item
    cost = guide.get("estimated_cost", {})
    if isinstance(cost, dict) and cost:
        add_card_header(doc, "💰 예상 비용 요약")
        doc.add_paragraph(f"• 교재: {cost.get('books', 0):,}원")
        doc.add_paragraph(f"• 강의: {cost.get('courses', 0):,}원")
        doc.add_paragraph(f"• 장비/기타: {cost.get('equipment', 0):,}원")
        doc.add_paragraph(f"• 총합: {cost.get('total', 0):,}원")
        divider = doc.add_paragraph("━" * 35)
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")
        if cost.get("breakdown"):
            doc.add_paragraph("")
            doc.add_paragraph("세부 비용", style='Heading 2')
            for item in cost["breakdown"]:
                line = doc.add_paragraph(style="List Bullet")
                line.add_run(f"{item.get('name')}: ").bold = True
                line.add_run(f"{item.get('average_price', 0):,}{item.get('currency', '원')}")
                if item.get("sources"):
                    src = item["sources"][0]
                    doc.add_paragraph(
                        f"   • 출처: {src.get('title', '')} ({src.get('url', '')})",
                        style="List Bullet 2"
                    )

            recommendations = select_recommended_items(cost["breakdown"], guide.get("category", ""))
            if recommendations:
                add_card_header(doc, "✨ 추천 Item")
                for idx, rec in enumerate(recommendations, 1):
                    para = doc.add_paragraph(f"{idx}. {rec.get('name', '추천 아이템')}")
                    para.runs[0].bold = True
                    para.add_run(f" - 약 {rec.get('average_price', 0):,}{rec.get('currency', '원')}")
                    if rec.get("sources"):
                        src = rec["sources"][0]
                        reference = doc.add_paragraph(style="List Bullet 2")
                        reference.add_run("참고: ")
                        add_hyperlink(reference, src.get("title", ""), src.get("url", ""))

        doc.add_paragraph("")

    # Step 카드
    for step in guide.get("steps", []):
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("")
        step_num = step.get("step_number", 0)
        step_title = clean_text(step.get("title", f"단계 {step_num}"))
        add_card_header(doc, f"🔵 Step {step_num}: {step_title}")

        period = doc.add_paragraph()
        period.add_run("📅 기간: ").bold = True
        period.add_run(f"{step.get('start_date', 'N/A')} ~ {step.get('end_date', 'N/A')} ({step.get('duration_days', 'N/A')}일)")

        # 학습 내용
        contents = normalize_learning_items(step.get("learning_content", []))
        if contents:
            doc.add_paragraph("📚 학습 내용").bold = True
            for item in contents:
                doc.add_paragraph(clean_text(item), style="List Bullet")
            doc.add_paragraph("")

        # 참고 사이트
        sites = step.get("recommended_sites", [])
        if sites:
            doc.add_paragraph("🌐 참고 자료").bold = True
            for site in sites:
                name = clean_text(site.get("name", "")) if isinstance(site, dict) else clean_text(site)
                url = site.get("url", "") if isinstance(site, dict) else ""
                para = doc.add_paragraph()
                para.add_run("🔗 ")
                if url:
                    add_hyperlink(para, name or url, url)
                else:
                    para.add_run(name)
            doc.add_paragraph("")

        # 투두리스트
        todos = normalize_learning_items(step.get("todos", []))
        if todos:
            doc.add_paragraph("📝 To-do List").bold = True
            for todo in todos:
                doc.add_paragraph(f"☐ {clean_text(todo)}")

        doc.add_paragraph("")

    # 후기 카드
    if guide.get("reviews_summary"):
        add_card_header(doc, "📌 학습 후기 요약")
        doc.add_paragraph(clean_text(guide["reviews_summary"]))

    doc.save(filename)
    print(f"✅ 워드 파일이 저장되었습니다: {filename}")
    return filename


def set_document_style(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Malgun Gothic"
    font.size = Pt(11)


